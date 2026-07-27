"""Ingest material into a spawn's knowledge base: extract text (files), strip
<private> blocks, chunk, and store to knowledge_chunks + FTS5."""
from __future__ import annotations

import io
import logging
import re

from sqlalchemy import text as sa_text

from arslan.core.chunking import chunk_text
from server.db import session as db_session
from server.db.models import KnowledgeChunk
from server.services.llm_factory import build_adapter
from server.services.prompts.kb_compress import COMPRESS_SYSTEM

logger = logging.getLogger(__name__)

_PRIVATE_RE = re.compile(r"<private>.*?</private>", re.DOTALL | re.IGNORECASE)
_OCR_MIN_CHARS = 20
_IMAGE_EXT_RE = re.compile(r"\.(png|jpe?g|webp|gif)$", re.IGNORECASE)


def _strip_private(text: str) -> str:
    return _PRIVATE_RE.sub("", text or "")


def _pdf_text_layer(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _ocr_pdf(data: bytes) -> str:
    """OCR a scanned PDF (no text layer). Best-effort; returns '' on any failure.

    Rasterizes with pypdfium2 (BSD-3-Clause / Apache-2.0). It replaced pymupdf
    in S4.3-a: pymupdf is AGPL-3.0, and shipping it inside the distributed .dmg
    would pull AGPL onto the whole app. See pyproject.toml's `ocr` extra.

    NOT INSTALLED IN THE PACKAGED APP. Both this import and pytesseract live in
    the opt-in `ocr` extra, so in a released build the ImportError below is the
    NORMAL path and scanned PDFs yield no text. The fix is the vision round
    (BYOK image input), not re-adding tesseract —
    see docs/specs/2026-07-26-s4.3a-packaging.md §9.
    """
    try:
        import pypdfium2 as pdfium
        import pytesseract
        out = []
        doc = pdfium.PdfDocument(data)
        try:
            for page in doc:
                # scale is a multiple of PDF user space (72 units/inch); the
                # 200 dpi target is what tesseract wants for body text.
                bitmap = page.render(scale=200 / 72)
                # Normalize the mode. MEASURED on pypdfium2 5.12.1: to_pil()
                # already hands back "RGB" for this render call, so today this
                # convert is a no-op — it is kept as a cheap guarantee that
                # pytesseract never sees RGBA/L from a future version or a
                # transparency-carrying render. Pinned by
                # test_whatever_the_rasterizer_returns_is_normalized_to_rgb,
                # which stubs to_pil() to RGBA because asserting the mode of a
                # REAL render cannot tell the two apart.
                img = bitmap.to_pil().convert("RGB")
                try:
                    out.append(pytesseract.image_to_string(img, lang="chi_sim+eng"))
                except Exception:  # noqa: BLE001 — language pack missing → English only
                    out.append(pytesseract.image_to_string(img))
        finally:
            doc.close()
        return "\n".join(out)
    except Exception as exc:  # noqa: BLE001 — OCR is best-effort
        logger.warning("OCR failed: %s", exc)
        return ""


# ── vision (S4.2-e) ────────────────────────────────────────────────────────
# What the model is asked to produce for a fed image, and how the result is
# labelled so a retrieval hit can never pass as text taken FROM the picture.

_DESCRIBE_SYSTEM = (
    "You are looking at an image the user filed into their knowledge base. "
    "Describe it so it can be found later and reasoned about: what it shows, "
    "and — if it contains text, tables or a UI — transcribe that content "
    "faithfully. Be concrete. Do not speculate about anything not visible."
)

# Stable and NOT localised, exactly like FACT_CATEGORIES: this is stored data.
# A translated marker would make the same picture read differently depending on
# which language the app happened to be in when it was fed.
_DESCRIBED_SUFFIX = " (image description)"

# How many pages of a scanned PDF may be sent to the model.
#
# NOT a round number picked for comfort — measured, because "there is a cap"
# without a number is not a cap. Each page is rasterised to the same 1568px long
# edge as any other image, i.e. ~1109x1568 for A4 portrait, which prices at
# ~2,319 tokens (Anthropic's pixels/750) or ~2,125 (OpenAI's 512px tiles).
#
#     pages   tokens    deepseek-flash   deepseek-pro   sonnet
#        20   46,371        $0.006          $0.020      $0.14
#        36   83,468        $0.012          $0.037      $0.25
#        50  115,927        $0.016          $0.051      $0.35
#       100  231,855        $0.032          $0.102      $0.70
#
# So money is NOT the binding constraint — the CONTEXT WINDOW is. At 36 pages
# (~83k) a 128k-context model still has ~45k left for the system prompt, the
# history and its own answer. At 50 pages (~116k) any history at all overflows
# it. The true ceiling sits near 40; 36 is the safe side of it.
VISION_PDF_MAX_PAGES = 36


def pdf_page_plan(total_pages: int) -> tuple[int, str]:
    """(pages to send, note to show the user).

    The note is EMPTY when nothing was skipped, and names both numbers when
    something was — a vague "some pages were skipped" is what lets someone act
    on 36 pages of a 137-page contract believing they had all of it."""
    take = min(max(total_pages, 0), VISION_PDF_MAX_PAGES)
    if total_pages > VISION_PDF_MAX_PAGES:
        return take, (f"This document has {total_pages} pages; the first "
                      f"{take} were read. The rest were not.")
    return take, ""


def described_source(filename: str) -> str:
    """Provenance marker for a model-described image.

    It rides on `source`, not as a text prefix, because a description long
    enough to be chunked would carry a prefix on the FIRST chunk only, whereas
    `source` is stored on EVERY chunk — and `source` is what knowledge_block
    renders into the model's view ("[source] text") and what the recall
    executor returns."""
    return f"{filename}{_DESCRIBED_SUFFIX}"


async def describe_image(data: bytes, mime_type: str) -> str:
    """Ask the configured model to describe an image. Raises on failure — the
    caller turns that into "stored nothing", which the UI reports honestly."""
    import base64

    from server.services.llm_factory import build_adapter

    adapter = await build_adapter(role="converse")
    blocks = [
        {"type": "text", "text": "Describe this image for a knowledge base."},
        {"type": "image", "mime_type": mime_type or "image/png",
         "data": base64.b64encode(data).decode()},
    ]
    resp = await adapter.chat(system=_DESCRIBE_SYSTEM, user=blocks)
    text = (resp.content or "").strip()
    if not text:
        raise RuntimeError("the model returned no description")
    return text


def _ocr_image(data: bytes) -> str:
    """OCR a standalone image attachment. Best-effort; returns '' on any failure
    (undecodable bytes, missing tesseract) or when no text is found. convert("RGB")
    normalizes RGBA/palette modes and takes an animated GIF's first frame."""
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(io.BytesIO(data)).convert("RGB")
        try:
            text = pytesseract.image_to_string(img, lang="chi_sim+eng")
        except Exception:  # noqa: BLE001 — language pack missing → English only
            text = pytesseract.image_to_string(img)
        return text if text.strip() else ""
    except Exception as exc:  # noqa: BLE001 — OCR is best-effort
        logger.warning("image OCR failed: %s", exc)
        return ""


def _extract_file(filename: str, data: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".txt") or name.endswith(".md"):
        return data.decode("utf-8", errors="replace")
    if name.endswith(".pdf"):
        text = _pdf_text_layer(data)
        if len(text.strip()) < _OCR_MIN_CHARS:
            try:
                ocr = _ocr_pdf(data)
            except Exception as exc:  # noqa: BLE001 — defense in depth (stubbed _ocr_pdf may raise)
                logger.warning("OCR fallback errored: %s", exc)
                ocr = ""
            if ocr.strip():
                return ocr
        return text
    if name.endswith(".docx"):
        import docx  # python-docx
        document = docx.Document(io.BytesIO(data))
        return "\n".join(p.text for p in document.paragraphs)
    if name.endswith((".html", ".htm")):
        import lxml.html  # available in the venv
        try:
            return lxml.html.fromstring(data).text_content()
        except Exception:  # noqa: BLE001 — bad/empty HTML → no text (caller surfaces "no text")
            return ""
    if _IMAGE_EXT_RE.search(name):
        # No text found / OCR unavailable → '' (caller surfaces "no text", never 500).
        return _ocr_image(data)
    raise ValueError(f"unsupported file type: {filename}")


async def _compress(text: str) -> str:
    """LLM-clean text; best-effort — returns original on failure/empty."""
    try:
        adapter = await build_adapter(role="converse")
        resp = await adapter.chat(system=COMPRESS_SYSTEM, user=text)
        cleaned = (resp.content or "").strip()
        return cleaned if cleaned else text
    except Exception as exc:  # noqa: BLE001
        logger.warning("kb compress failed, using original: %s", exc)
        return text


async def ingest_text(spawn_id: int | None, source: str, text: str, *,
                      collection_id: int | None = None, compress: bool = False) -> int:
    """Strip private blocks, chunk, store into EXACTLY ONE of a spawn's private
    well (spawn_id) or a shared collection (collection_id), then best-effort
    embed the new chunks. Returns number of chunks stored."""
    if (spawn_id is None) == (collection_id is None):
        raise ValueError("provide exactly one of spawn_id / collection_id")
    cleaned = _strip_private(text)
    if compress:
        cleaned = await _compress(cleaned)
    chunks = chunk_text(cleaned)
    if not chunks:
        return 0
    ids_texts: list[tuple[int, str]] = []
    async with db_session.AsyncSessionLocal() as db:
        for i, chunk in enumerate(chunks):
            row = KnowledgeChunk(spawn_id=spawn_id, collection_id=collection_id,
                                 source=source, chunk_index=i, text=chunk)
            db.add(row)
            await db.flush()  # populate row.id for the FTS rowid
            # Defensive: clear any stale FTS row at this rowid first. knowledge_chunks
            # ids restart from 1 when the table is emptied, so a leftover FTS orphan
            # (from a crash / partial delete / migration edge) would otherwise collide
            # and 500 every ingest. Delete-then-insert keeps the FTS mirror consistent.
            await db.execute(
                sa_text("DELETE FROM knowledge_chunks_fts WHERE rowid = :rid"),
                {"rid": row.id},
            )
            await db.execute(
                sa_text("INSERT INTO knowledge_chunks_fts (rowid, text) VALUES (:rid, :t)"),
                {"rid": row.id, "t": chunk},
            )
            ids_texts.append((row.id, chunk))
        await db.commit()
    await _embed_new_chunks(ids_texts)
    return len(chunks)


async def _embed_new_chunks(ids_texts: list[tuple[int, str]]) -> None:
    """Vectorize freshly stored chunks. Best-effort: any failure leaves the
    embedding NULL and the chunks retrievable via FTS only. All chunks from one
    ingest are sent as a single un-chunked provider batch, so an oversized
    input can fail the whole batch at once; recovery relies on the
    embed_missing() backfill picking up the resulting NULL rows later."""
    if not ids_texts:
        return
    try:
        from server.services import embedding_service
        provider = await embedding_service.active_provider()
        if provider is None:
            return
        vecs = await provider.embed([t for _, t in ids_texts])
        async with db_session.AsyncSessionLocal() as db:
            for (cid, _), vec in zip(ids_texts, vecs):
                await db.execute(
                    sa_text("UPDATE knowledge_chunks SET embedding = :b, embedding_model = :m "
                            "WHERE id = :id"),
                    {"b": embedding_service.vec_to_blob(vec), "m": provider.model_id, "id": cid},
                )
            await db.commit()
    except Exception as exc:  # noqa: BLE001 — embedding is never fatal
        logger.warning("chunk embedding failed (non-fatal, FTS-only): %s", exc)


async def ingest_url(spawn_id: int | None, url: str, *,
                     collection_id: int | None = None, compress: bool = False) -> int:
    """Fetch + extract a web page via the SSRF-guarded WebExtractExecutor (per-hop
    host revalidation + private-IP block — NEVER a raw httpx request), then ingest.
    Raises ValueError on fetch failure / private-address rejection."""
    from server.registry.executors import EXECUTORS
    res = await EXECUTORS["web_extract"].execute({"url": url})
    if not res.get("ok"):
        raise ValueError(res.get("error") or "fetch failed")
    return await ingest_text(spawn_id, url, res.get("text", ""),
                             collection_id=collection_id, compress=compress)


async def ingest_file(spawn_id: int | None, filename: str, data: bytes, *,
                      collection_id: int | None = None, compress: bool = False) -> int:
    """Extract text from a supported file then ingest. Raises ValueError on
    unsupported extension; extraction errors propagate (API maps to 400)."""
    if _IMAGE_EXT_RE.search(filename):
        # The MODEL reads the picture (vision round). A failure here means the
        # image could not be read at all — surface it as zero chunks so the UI
        # says so, rather than storing an empty "success".
        try:
            described = await describe_image(data, "image/png")
        except Exception as exc:  # noqa: BLE001 — never fatal; honest zero
            logger.warning("image description failed (stored nothing): %s", exc)
            return 0
        return await ingest_text(spawn_id, described_source(filename), described,
                                 collection_id=collection_id, compress=compress)
    extracted = _extract_file(filename, data)
    return await ingest_text(spawn_id, filename, extracted,
                             collection_id=collection_id, compress=compress)
